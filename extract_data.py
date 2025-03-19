import json
from docx import Document

print("🚀 Script started...")  

# Load the Word document
file_path = "data_001.docx"
doc = Document(file_path)

print("✅ Document loaded successfully!")  

# Initialize structured JSON storage
structured_data = {
    "document_title": "Data001",
    "metadata": "Created July 2022",
    "table_of_contents": [],
    "sections": []
}

current_section = None  # Store current section
last_table_heading = None  # Track last detected table heading
table_description = []  # Store text before a table
found_toc = False  # Track when Table of Contents starts

# Convert document paragraphs to a list for easier indexing
paragraphs = list(doc.paragraphs)

for i, para in enumerate(paragraphs):
    text = para.text.strip()

    if not text:
        continue  # Skip empty paragraphs

    print(f"🔹 Processing Paragraph {i+1}: {text[:50]}")  # Debugging

    # Detect Table of Contents
    if text.lower() == "table of contents":
        found_toc = True
        continue

    if found_toc:
        structured_data["table_of_contents"].append({"title": text, "linked_section": text})
        continue

    # Detect Headings
    if para.style.name.startswith("Heading"):
        # Store previous section before moving to a new one
        if current_section:
            structured_data["sections"].append(current_section)

        # Start new section
        current_section = {"header": text, "content": [], "tables": []}
        last_table_heading = None
        table_description = []

    elif "Table" in text or text.isupper():  # Detect table headings
        last_table_heading = text

    elif last_table_heading:  # If text appears after a heading but before a table, it's a description
        table_description.append(text)

    else:
        if current_section:
            current_section["content"].append(text)

# Save last section
if current_section:
    structured_data["sections"].append(current_section)

# Extract tables and assign them to the correct heading
for i, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)

    # Find the closest heading before the table
    if i > 0:
        for j in range(i - 1, -1, -1):
            if paragraphs[j].text.strip():
                last_table_heading = paragraphs[j].text.strip()
                break

    # Store table with the correct name and description
    if structured_data["sections"]:
        structured_data["sections"][-1]["tables"].append({
            "name": last_table_heading if last_table_heading else "Unnamed Table",
            "description": " ".join(table_description),
            "data": table_data
        })
        last_table_heading = None
        table_description = []
    else:
        structured_data["sections"].append({
            "header": "Tables",
            "content": [],
            "tables": [{"name": last_table_heading if last_table_heading else "Unnamed Table", "description": " ".join(table_description), "data": table_data}]
        })

# Save to JSON
output_file = "output.json"
with open(output_file, "w", encoding="utf-8") as json_file:
    json.dump(structured_data, json_file, indent=4, ensure_ascii=False)

print(f"✅ Data saved successfully in '{output_file}'!")
